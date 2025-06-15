import os
import sys
from typing import List, Dict
from pathlib import Path

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.schema import Document

# プロジェクト設定
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# 設定読み込み（緊急修正）
try:
    from src.config.settings import settings
    OPENAI_API_KEY = settings.OPENAI_API_KEY
    FAISS_INDEX_DIR = str(settings.FAISS_INDEX_DIR)
except ImportError:
    # フォールバック - 直接.envから読み込み
    from dotenv import load_dotenv
    load_dotenv()
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    FAISS_INDEX_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "data", "faiss_index")
    
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found in environment variables or .env file")

# sitecustomize.py でのパッチ適用確認
try:
    import sitecustomize
    if hasattr(sitecustomize, 'OPENAI_PROXIES_PATCH_APPLIED'):
        print("✅ OpenAI proxies patch confirmed from sitecustomize.py")
    else:
        print("⚠️ sitecustomize.py found but patch flag not detected")
except ImportError:
    print("⚠️ sitecustomize.py not found - patch may not be applied")

class QAGenerator:
    def __init__(self):
        # OpenAI API キーを環境変数として設定
        os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
        
        self.embeddings = OpenAIEmbeddings()
        self.llm = ChatOpenAI(
            temperature=0.7,
            model_name="gpt-4o"
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    def process_document(self, file_path: str, lecture_id: int) -> bool:
        """
        ドキュメントを処理してFAISSインデックスを作成
        """
        try:
            # ファイル読み込み
            content = self._read_file(file_path)
            if not content:
                print(f"Error: Could not read file {file_path}")
                return False
            
            # テキスト分割
            documents = self.text_splitter.create_documents([content])
            
            # メタデータを追加
            for doc in documents:
                doc.metadata = {
                    "lecture_id": lecture_id,
                    "source": file_path
                }
            
            # FAISSインデックス作成
            vectorstore = FAISS.from_documents(documents, self.embeddings)
            
            # インデックス保存
            index_path = os.path.join(FAISS_INDEX_DIR, f"lecture_{lecture_id}")
            os.makedirs(index_path, exist_ok=True)
            vectorstore.save_local(index_path)
            
            print(f"Successfully processed document for lecture {lecture_id}")
            return True
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return False
    
    def generate_qa(self, lecture_id: int, difficulty: str, num_questions: int, question_types: List[str] = None) -> List[Dict[str, str]]:
        """
        指定された講義からQ&Aを生成（タイムアウト付き）
        """
        import time
        start_time = time.time()
        timeout_seconds = 120  # 2分でタイムアウト
        
        try:
            # デフォルトの質問タイプ
            if question_types is None:
                question_types = ["multiple_choice", "short_answer"]
            
            # FAISSインデックス読み込み
            index_path = os.path.join(FAISS_INDEX_DIR, f"lecture_{lecture_id}")
            if not os.path.exists(index_path):
                print(f"Error: FAISS index not found for lecture {lecture_id}")
                return []
            
            vectorstore = FAISS.load_local(
                index_path, 
                self.embeddings,
                allow_dangerous_deserialization=True
            )
            
            # Q&A生成（重複防止機能付き）
            generated_qas = []
            generated_questions = set()  # 重複チェック用
            max_attempts = min(num_questions * 2, 10)  # 最大試行回数を制限（安全のため）
            attempts = 0
            
            while len(generated_qas) < num_questions and attempts < max_attempts:
                # タイムアウトチェック
                if time.time() - start_time > timeout_seconds:
                    print(f"Timeout reached ({timeout_seconds}s). Generated {len(generated_qas)}/{num_questions} questions.")
                    break
                
                try:
                    attempts += 1
                    current_index = len(generated_qas)
                    
                    # 質問タイプをローテーション
                    question_type = question_types[current_index % len(question_types)]
                    
                    # RetrievalQA チェーン作成（質問タイプ別）
                    qa_chain = RetrievalQA.from_chain_type(
                        llm=self.llm,
                        chain_type="stuff",
                        retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
                        return_source_documents=False,
                        chain_type_kwargs={"prompt": self._get_qa_prompt(difficulty, question_type)}
                    )
                    
                    # 質問生成のためのクエリ（多様性を高めるため）
                    variety_keywords = ["基本的な", "重要な", "具体的な", "実践的な", "理論的な"]
                    variety_keyword = variety_keywords[attempts % len(variety_keywords)]
                    
                    query = f"講義内容に基づいて{variety_keyword}{difficulty}レベルの{self._get_question_type_name(question_type)}を1つ作成してください。これまでに作成された質問とは異なる内容で、質問番号: {current_index+1}"
                    
                    result = qa_chain.invoke({"query": query})
                    qa_text = result["result"]
                    
                    # Q&Aを分離
                    qa_pair = self._parse_qa_response(qa_text, difficulty, question_type)
                    if qa_pair and qa_pair.get("question"):
                        # 重複チェック（質問の最初の30文字で判定、より厳密に）
                        question_key = qa_pair["question"][:30].strip().lower().replace(" ", "").replace("　", "")
                        if question_key not in generated_questions and len(question_key) > 5:
                            qa_pair['question_type'] = question_type
                            generated_qas.append(qa_pair)
                            generated_questions.add(question_key)
                            print(f"Generated unique QA {len(generated_qas)}/{num_questions} (type: {question_type})")
                        else:
                            print(f"Duplicate question detected, retrying... (attempt {attempts})")
                        
                except Exception as e:
                    print(f"Error generating QA (attempt {attempts}): {str(e)}")
                    continue
            
            if len(generated_qas) < num_questions:
                print(f"Warning: Only generated {len(generated_qas)}/{num_questions} unique questions after {attempts} attempts")
            
            return generated_qas
            
        except Exception as e:
            print(f"Error in generate_qa: {str(e)}")
            return []
    
    def _read_file(self, file_path: str) -> str:
        """
        ファイルを読み込む（拡張子に応じて処理を分岐）
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_extension == '.pdf':
                return self._read_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._read_word_document(file_path)
            else:
                # テキストファイルとして読み込み試行
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            print(f"Error reading file {file_path}: {str(e)}")
            return ""
    
    def _read_pdf(self, file_path: str) -> str:
        """
        PDFファイルを読み込む（OCRフォールバック付き）
        """
        try:
            import PyPDF2
            
            # まずPyPDF2で試行
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
                # テキストが十分に抽出できた場合
                if len(text.strip()) > 100:  # 100文字以上あれば成功とみなす
                    print(f"Successfully read PDF with PyPDF2: {len(text)} characters")
                    return text
                else:
                    print("PyPDF2 extracted insufficient text, trying OCR fallback...")
                    return self._read_pdf_with_ocr(file_path)
                
        except Exception as e:
            print(f"PyPDF2 failed: {str(e)}, trying OCR fallback...")
            return self._read_pdf_with_ocr(file_path)
    
    def _read_pdf_with_ocr(self, file_path: str) -> str:
        """
        OCRを使用してPDFからテキストを抽出
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image
            
            print("Starting OCR processing...")
            
            # PDFを画像に変換
            pages = convert_from_path(file_path, dpi=200)
            text = ""
            
            for i, page in enumerate(pages):
                print(f"Processing page {i+1}/{len(pages)} with OCR...")
                
                # OCRでテキスト抽出（日本語対応）
                page_text = pytesseract.image_to_string(page, lang='jpn+eng')
                text += page_text + "\n"
            
            if len(text.strip()) > 50:  # 50文字以上あれば成功
                print(f"Successfully read PDF with OCR: {len(text)} characters")
                return text
            else:
                print("OCR also failed to extract sufficient text")
                return ""
                
        except Exception as e:
            print(f"OCR processing failed: {str(e)}")
            # Tesseractがインストールされていない場合の対処
            if "tesseract" in str(e).lower():
                print("Tesseract not installed. Please install: sudo apt-get install tesseract-ocr tesseract-ocr-jpn")
            return ""
    
    def _read_word_document(self, file_path: str) -> str:
        """
        Word文書（DOCX/DOC）を読み込む
        """
        try:
            from docx import Document
            
            # DOCXファイルの読み込み
            if file_path.lower().endswith('.docx'):
                doc = Document(file_path)
                text = ""
                
                # 段落のテキストを抽出
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # テーブルのテキストも抽出
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                
                print(f"Successfully read DOCX: {len(text)} characters")
                return text
            
            # DOCファイルの場合（python-docxはDOCXのみ対応）
            elif file_path.lower().endswith('.doc'):
                print("DOC format not directly supported. Please convert to DOCX or use LibreOffice conversion.")
                # 将来的にはpython-docx2txtやlibreoffice経由での変換を実装可能
                return ""
                
        except Exception as e:
            print(f"Error reading Word document {file_path}: {str(e)}")
            return ""
    
    def _get_question_type_name(self, question_type: str) -> str:
        """質問タイプの日本語名を取得"""
        type_names = {
            "multiple_choice": "選択問題",
            "short_answer": "短答問題", 
            "essay": "記述問題"
        }
        return type_names.get(question_type, "質問")
    
    def _get_qa_prompt(self, difficulty: str, question_type: str = "multiple_choice") -> PromptTemplate:
        """
        難易度と質問タイプに応じたプロンプトテンプレートを取得
        """
        difficulty_instructions = {
            "easy": "基本的な概念や定義に関する簡単な",
            "medium": "概念の理解や応用に関する中程度の",
            "hard": "深い理解や批判的思考を要する難しい"
        }
        
        type_instructions = {
            "multiple_choice": """4択の選択問題を作成してください。
以下の形式で回答してください：
質問: [ここに質問を記載]
A) [選択肢1]
B) [選択肢2] 
C) [選択肢3]
D) [選択肢4]
正解: [A/B/C/Dのいずれか]
解説: [正解の理由を簡潔に説明]""",
            
            "short_answer": """短答問題を作成してください。
以下の形式で回答してください：
質問: [ここに質問を記載]
回答: [簡潔な回答（1-2文程度）]
解説: [回答の補足説明]""",
            
            "essay": """記述問題を作成してください。
以下の形式で回答してください：
質問: [ここに質問を記載]
回答: [詳細な回答（3-5文程度）]
評価ポイント: [回答で重視すべき要素]"""
        }
        
        template = f"""
以下の文脈に基づいて、{difficulty_instructions.get(difficulty, "適切な")}{self._get_question_type_name(question_type)}を1つ作成してください。

文脈: {{context}}

要求: {{question}}

{type_instructions.get(question_type, type_instructions["multiple_choice"])}

質問と回答は明確で、文脈に基づいた内容にしてください。
"""
        
        return PromptTemplate(
            template=template,
            input_variables=["context", "question"]
        )
    
    def _parse_qa_response(self, response: str, difficulty: str, question_type: str = "multiple_choice") -> Dict[str, str]:
        """
        LLMの応答からQ&Aペアを抽出（質問タイプ別）
        """
        try:
            lines = response.strip().split('\n')
            question = ""
            answer = ""
            choices = []
            correct_answer = ""
            explanation = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('質問:') or line.startswith('Q:'):
                    question = line.split(':', 1)[1].strip()
                elif line.startswith('回答:') or line.startswith('A:'):
                    answer = line.split(':', 1)[1].strip()
                elif line.startswith('正解:'):
                    correct_answer = line.split(':', 1)[1].strip()
                elif line.startswith('解説:'):
                    explanation = line.split(':', 1)[1].strip()
                elif question_type == "multiple_choice" and line.startswith(('A)', 'B)', 'C)', 'D)')):
                    choices.append(line)
            
            # 選択問題の場合、選択肢と正解を含む完全な回答を作成
            if question_type == "multiple_choice" and choices:
                full_answer = "\n".join(choices)
                if correct_answer:
                    full_answer += f"\n\n正解: {correct_answer}"
                if explanation:
                    full_answer += f"\n解説: {explanation}"
                answer = full_answer
            elif not answer:
                answer = "回答が見つかりませんでした。"
            
            if question:
                result = {
                    "question": question,
                    "answer": answer,
                    "difficulty": difficulty,
                    "question_type": question_type
                }
                
                # 選択問題の場合、追加情報を保存
                if question_type == "multiple_choice":
                    result["choices"] = choices
                    result["correct_answer"] = correct_answer
                    result["explanation"] = explanation
                
                return result
            else:
                # フォールバック - 応答全体を使用
                return {
                    "question": response.strip()[:200] + "..." if len(response) > 200 else response.strip(),
                    "answer": "この質問については講義資料を参照してください。",
                    "difficulty": difficulty,
                    "question_type": question_type
                }
                
        except Exception as e:
            print(f"Error parsing QA response: {str(e)}")
            return None


# シングルトンインスタンス
qa_generator = QAGenerator() 